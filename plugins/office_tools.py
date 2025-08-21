from PyQt5.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QLineEdit, 
    QPushButton, QTextEdit, QComboBox, QGroupBox, QCheckBox,
    QProgressBar, QMessageBox, QFileDialog,
    QTableWidget, QTableWidgetItem, QHeaderView
)
from PyQt5.QtCore import Qt, QThread, pyqtSignal, QMutex, QMutexLocker
from PyQt5.QtGui import QFont, QTextCursor
from plugins.base_plugin import BasePlugin
import os  # 添加os模块的导入
import fitz  # PyMuPDF库，用于读取PDF文件
from docx import Document  # python-docx库，用于创建Word文档
from docx.shared import Inches
import pandas as pd  # 用于处理Excel和CSV文件
import comtypes.client  # 用于Word转PDF（Windows系统）
import tempfile
import shutil


class WordToPDFPlugin(BasePlugin):
    def __init__(self):
        super().__init__()
        self.name = "Word转PDF"
        self.description = "将Word文档转换为PDF格式"
        self.category = "办公工具"
        self.icon = None

    def get_widget(self):
        """返回Word转PDF插件界面"""
        return WordToPDFWidget()


class PDFToWordPlugin(BasePlugin):
    def __init__(self):
        super().__init__()
        self.name = "PDF转Word"
        self.description = "将PDF文档转换为Word格式"
        self.category = "办公工具"
        self.icon = None

    def get_widget(self):
        """返回PDF转Word插件界面"""
        return PDFToWordWidget()


class ExcelExtractPlugin(BasePlugin):
    def __init__(self):
        super().__init__()
        self.name = "Excel数据处理"
        self.description = "提取和处理Excel表格数据"
        self.category = "办公工具"
        self.icon = None

    def get_widget(self):
        """返回Excel数据处理插件界面"""
        return ExcelExtractWidget()


# Word转PDF插件界面
class WordToPDFWidget(QWidget):
    def __init__(self):
        super().__init__()
        self.init_ui()
        self.conversion_thread = None
    
    def __del__(self):
        # 应用程序关闭时释放Word实例
        WordToPDFThread.release_word_app()
    
    def init_ui(self):
        layout = QVBoxLayout()
        
        # 标题
        title_label = QLabel("Word转PDF工具")
        title_font = QFont("Arial", 12, QFont.Bold)
        title_label.setFont(title_font)
        title_label.setAlignment(Qt.AlignCenter)
        title_label.setStyleSheet("color: #2c3e50; padding: 10px;")
        layout.addWidget(title_label)
        
        # 文件选择区域
        file_group = QGroupBox("文件选择")
        file_layout = QVBoxLayout()
        
        input_layout = QHBoxLayout()
        self.word_input = QLineEdit()
        self.word_input.setPlaceholderText("选择Word文档...")
        self.word_browse_btn = QPushButton("浏览...")
        self.word_browse_btn.clicked.connect(self.browse_word_file)
        
        input_layout.addWidget(QLabel("Word文件:"))
        input_layout.addWidget(self.word_input)
        input_layout.addWidget(self.word_browse_btn)
        
        output_layout = QHBoxLayout()
        self.pdf_output = QLineEdit()
        self.pdf_output.setPlaceholderText("选择PDF输出位置...")
        self.pdf_browse_btn = QPushButton("浏览...")
        self.pdf_browse_btn.clicked.connect(self.browse_pdf_output)
        
        output_layout.addWidget(QLabel("PDF输出:"))
        output_layout.addWidget(self.pdf_output)
        output_layout.addWidget(self.pdf_browse_btn)
        
        file_layout.addLayout(input_layout)
        file_layout.addLayout(output_layout)
        file_group.setLayout(file_layout)
        
        # 转换选项
        options_group = QGroupBox("转换选项")
        options_layout = QVBoxLayout()
        
        self.keep_layout_check = QCheckBox("保持原始布局")
        self.keep_layout_check.setChecked(True)
        
        self.optimize_pdf_check = QCheckBox("优化PDF文件大小")
        self.optimize_pdf_check.setChecked(False)
        
        self.fast_mode_check = QCheckBox("快速转换模式 (推荐)")
        self.fast_mode_check.setChecked(True)  # 默认启用快速模式
        
        options_layout.addWidget(self.keep_layout_check)
        options_layout.addWidget(self.optimize_pdf_check)
        options_layout.addWidget(self.fast_mode_check)
        options_group.setLayout(options_layout)
        
        # 控制按钮
        control_layout = QHBoxLayout()
        self.convert_btn = QPushButton("开始转换")
        self.convert_btn.clicked.connect(self.convert_word_to_pdf)
        
        self.cancel_btn = QPushButton("取消")
        self.cancel_btn.clicked.connect(self.cancel_conversion)
        self.cancel_btn.setEnabled(False)
        
        control_layout.addStretch()
        control_layout.addWidget(self.convert_btn)
        control_layout.addWidget(self.cancel_btn)
        
        # 进度条
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        
        # 日志区域
        self.log_area = QTextEdit()
        self.log_area.setReadOnly(True)
        self.log_area.setMaximumHeight(100)
        
        layout.addWidget(file_group)
        layout.addWidget(options_group)
        layout.addLayout(control_layout)
        layout.addWidget(self.progress_bar)
        layout.addWidget(self.log_area)
        
        self.setLayout(layout)
    
    def browse_word_file(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self, "选择Word文档", "", "Word文档 (*.docx *.doc);;所有文件 (*)"
        )
        if file_path:
            self.word_input.setText(file_path)
            # 自动生成PDF输出路径
            base_name = os.path.splitext(file_path)[0]
            self.pdf_output.setText(base_name + ".pdf")
    
    def browse_pdf_output(self):
        file_path, _ = QFileDialog.getSaveFileName(
            self, "保存PDF文件", "", "PDF文件 (*.pdf);;所有文件 (*)"
        )
        if file_path:
            self.pdf_output.setText(file_path)
    
    def convert_word_to_pdf(self):
        """转换Word到PDF"""
        word_file = self.word_input.text().strip()
        pdf_file = self.pdf_output.text().strip()
        
        if not word_file or not pdf_file:
            QMessageBox.warning(self, "错误", "请选择输入和输出文件")
            return
        
        if not os.path.exists(word_file):
            QMessageBox.warning(self, "错误", "Word文件不存在")
            return
        
        # 显示进度条
        self.progress_bar.setVisible(True)
        self.progress_bar.setRange(0, 100)  # 设置确定进度范围
        self.progress_bar.setValue(0)
        
        # 记录日志
        self.log_area.append(f"开始转换: {os.path.basename(word_file)} -> {os.path.basename(pdf_file)}")
        
        # 更新按钮状态
        self.convert_btn.setEnabled(False)
        self.cancel_btn.setEnabled(True)
        
        # 创建一个新的线程来执行转换，避免界面卡顿
        self.conversion_thread = WordToPDFThread(
            word_file, pdf_file, 
            self.keep_layout_check.isChecked(),
            self.optimize_pdf_check.isChecked(),
            self.fast_mode_check.isChecked()
        )
        self.conversion_thread.finished.connect(self.finish_conversion_thread)
        self.conversion_thread.progress_updated.connect(self.update_progress)
        self.conversion_thread.start()
    
    def update_progress(self, value):
        """更新进度条"""
        self.progress_bar.setValue(value)

    def cancel_conversion(self):
        """取消转换"""
        if self.conversion_thread and self.conversion_thread.isRunning():
            self.conversion_thread.cancel()
            self.log_area.append("正在取消转换...")

    def finish_conversion_thread(self, success, message):
        """线程完成后的回调"""
        self.progress_bar.setVisible(False)
        self.convert_btn.setEnabled(True)
        self.cancel_btn.setEnabled(False)
        
        if success:
            self.log_area.append(f"转换完成: {message}")
            QMessageBox.information(self, "成功", f"文件已转换并保存为: {message}")
        else:
            self.log_area.append(f"转换失败: {message}")
            QMessageBox.critical(self, "失败", f"转换失败: {message}")


# PDF转Word插件界面
class PDFToWordWidget(QWidget):
    def __init__(self):
        super().__init__()
        self.init_ui()
    
    def init_ui(self):
        layout = QVBoxLayout()
        
        # 标题
        title_label = QLabel("PDF转Word工具")
        title_font = QFont("Arial", 12, QFont.Bold)
        title_label.setFont(title_font)
        title_label.setAlignment(Qt.AlignCenter)
        title_label.setStyleSheet("color: #2c3e50; padding: 10px;")
        layout.addWidget(title_label)
        
        # 文件选择区域
        file_group = QGroupBox("文件选择")
        file_layout = QVBoxLayout()
        
        input_layout = QHBoxLayout()
        self.pdf_input = QLineEdit()
        self.pdf_input.setPlaceholderText("选择PDF文档...")
        self.pdf_browse_btn = QPushButton("浏览...")
        self.pdf_browse_btn.clicked.connect(self.browse_pdf_file)
        
        input_layout.addWidget(QLabel("PDF文件:"))
        input_layout.addWidget(self.pdf_input)
        input_layout.addWidget(self.pdf_browse_btn)
        
        output_layout = QHBoxLayout()
        self.word_output = QLineEdit()
        self.word_output.setPlaceholderText("选择Word输出位置...")
        self.word_browse_btn = QPushButton("浏览...")
        self.word_browse_btn.clicked.connect(self.browse_word_output)
        
        output_layout.addWidget(QLabel("Word输出:"))
        output_layout.addWidget(self.word_output)
        output_layout.addWidget(self.word_browse_btn)
        
        file_layout.addLayout(input_layout)
        file_layout.addLayout(output_layout)
        file_group.setLayout(file_layout)
        
        # 转换选项
        options_group = QGroupBox("转换选项")
        options_layout = QVBoxLayout()
        
        self.extract_images_check = QCheckBox("提取图片")
        self.extract_images_check.setChecked(True)
        
        self.preserve_formatting_check = QCheckBox("保留格式")
        self.preserve_formatting_check.setChecked(True)
        
        options_layout.addWidget(self.extract_images_check)
        options_layout.addWidget(self.preserve_formatting_check)
        options_group.setLayout(options_layout)
        
        # 控制按钮
        control_layout = QHBoxLayout()
        self.convert_btn = QPushButton("开始转换")
        self.convert_btn.clicked.connect(self.convert_pdf_to_word)
        
        control_layout.addStretch()
        control_layout.addWidget(self.convert_btn)
        
        # 进度条
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        
        # 日志区域
        self.log_area = QTextEdit()
        self.log_area.setReadOnly(True)
        self.log_area.setMaximumHeight(100)
        
        layout.addWidget(file_group)
        layout.addWidget(options_group)
        layout.addLayout(control_layout)
        layout.addWidget(self.progress_bar)
        layout.addWidget(self.log_area)
        
        self.setLayout(layout)
    
    def browse_pdf_file(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self, "选择PDF文档", "", "PDF文件 (*.pdf);;所有文件 (*)"
        )
        if file_path:
            self.pdf_input.setText(file_path)
            # 自动生成Word输出路径
            base_name = os.path.splitext(file_path)[0]
            self.word_output.setText(base_name + ".docx")
    
    def browse_word_output(self):
        file_path, _ = QFileDialog.getSaveFileName(
            self, "保存Word文档", "", "Word文档 (*.docx);;所有文件 (*)"
        )
        if file_path:
            self.word_output.setText(file_path)
    
    def convert_pdf_to_word(self):
        """转换PDF到Word"""
        pdf_file = self.pdf_input.text().strip()
        word_file = self.word_output.text().strip()
        
        if not pdf_file or not word_file:
            QMessageBox.warning(self, "错误", "请选择输入和输出文件")
            return
        
        if not os.path.exists(pdf_file):
            QMessageBox.warning(self, "错误", "PDF文件不存在")
            return
        
        # 显示进度条
        self.progress_bar.setVisible(True)
        self.progress_bar.setRange(0, 0)  # 不确定进度
        
        # 记录日志
        self.log_area.append(f"开始转换: {os.path.basename(pdf_file)} -> {os.path.basename(word_file)}")
        
        # 创建一个新的线程来执行转换，避免界面卡顿
        self.conversion_thread = PDFToWordThread(pdf_file, word_file, 
                           self.extract_images_check.isChecked())
        self.conversion_thread.finished.connect(self.finish_conversion_thread)
        self.conversion_thread.start()
    
    def finish_conversion_thread(self, success, message):
        """线程完成后的回调"""
        self.progress_bar.setVisible(False)
        if success:
            self.log_area.append(f"转换完成: {message}")
            QMessageBox.information(self, "成功", f"文件已转换并保存为: {message}")
        else:
            self.log_area.append(f"转换失败: {message}")
            QMessageBox.critical(self, "失败", f"转换失败: {message}")


class ExcelExtractWidget(QWidget):
    def __init__(self):
        super().__init__()
        self.is_csv_file = False  # 添加标志来区分CSV文件
        self.init_ui()
    
    def init_ui(self):
        layout = QVBoxLayout()
        
        # 标题
        title_label = QLabel("Excel数据处理工具")
        title_font = QFont("Arial", 12, QFont.Bold)
        title_label.setFont(title_font)
        title_label.setAlignment(Qt.AlignCenter)
        title_label.setStyleSheet("color: #2c3e50; padding: 10px;")
        layout.addWidget(title_label)
        
        # 文件选择区域
        file_group = QGroupBox("文件选择")
        file_layout = QVBoxLayout()
        
        input_layout = QHBoxLayout()
        self.excel_input = QLineEdit()
        self.excel_input.setPlaceholderText("选择Excel文件...")
        self.excel_browse_btn = QPushButton("浏览...")
        self.excel_browse_btn.clicked.connect(self.browse_excel_file)
        
        input_layout.addWidget(QLabel("Excel文件:"))
        input_layout.addWidget(self.excel_input)
        input_layout.addWidget(self.excel_browse_btn)
        
        file_layout.addLayout(input_layout)
        file_group.setLayout(file_layout)
        
        # 提取选项
        extract_group = QGroupBox("提取选项")
        extract_layout = QVBoxLayout()
        
        # 工作表选择
        sheet_layout = QHBoxLayout()
        sheet_layout.addWidget(QLabel("选择工作表:"))
        self.sheet_combo = QComboBox()
        sheet_layout.addWidget(self.sheet_combo)
        sheet_layout.addStretch()
        
        # 提取范围
        range_layout = QHBoxLayout()
        range_layout.addWidget(QLabel("提取范围:"))
        self.range_start = QLineEdit()
        self.range_start.setPlaceholderText("起始单元格，如A1")
        self.range_end = QLineEdit()
        self.range_end.setPlaceholderText("结束单元格，如D10")
        range_layout.addWidget(self.range_start)
        range_layout.addWidget(QLabel("到"))
        range_layout.addWidget(self.range_end)
        
        # 提取模式
        mode_layout = QHBoxLayout()
        mode_layout.addWidget(QLabel("提取模式:"))
        self.mode_combo = QComboBox()
        self.mode_combo.addItems(["所有数据", "仅特定列", "仅特定行"])
        mode_layout.addWidget(self.mode_combo)
        
        extract_layout.addLayout(sheet_layout)
        extract_layout.addLayout(range_layout)
        extract_layout.addLayout(mode_layout)
        extract_group.setLayout(extract_layout)
        
        # 控制按钮
        control_layout = QHBoxLayout()
        self.extract_btn = QPushButton("提取数据")
        self.extract_btn.clicked.connect(self.extract_excel_data)
        
        control_layout.addStretch()
        control_layout.addWidget(self.extract_btn)
        
        # 结果显示
        result_group = QGroupBox("提取结果")
        result_layout = QVBoxLayout()
        
        self.result_table = QTableWidget()
        self.result_table.setAlternatingRowColors(True)
        self.result_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        
        result_layout.addWidget(self.result_table)
        result_group.setLayout(result_layout)
        
        # 导出按钮
        export_layout = QHBoxLayout()
        self.export_csv_btn = QPushButton("导出为CSV")
        self.export_csv_btn.clicked.connect(self.export_to_csv)
        self.export_excel_btn = QPushButton("导出为Excel")
        self.export_excel_btn.clicked.connect(self.export_to_excel)
        # 添加导出为TXT按钮
        self.export_txt_btn = QPushButton("导出为TXT")
        self.export_txt_btn.clicked.connect(self.export_to_txt)
        
        export_layout.addStretch()
        export_layout.addWidget(self.export_csv_btn)
        export_layout.addWidget(self.export_excel_btn)
        export_layout.addWidget(self.export_txt_btn)  # 添加到布局中
        
        layout.addWidget(file_group)
        layout.addWidget(extract_group)
        layout.addLayout(control_layout)
        layout.addWidget(result_group)
        layout.addLayout(export_layout)
        
        self.setLayout(layout)
    
    def browse_excel_file(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self, "选择数据文件", "", "数据文件 (*.xlsx *.xls *.csv);;Excel文件 (*.xlsx *.xls);;CSV文件 (*.csv);;所有文件 (*)"
        )
        if file_path:
            self.excel_input.setText(file_path)
            # 判断文件类型
            self.is_csv_file = file_path.lower().endswith('.csv')
            # 加载工作表列表或处理CSV
            self.load_sheet_list(file_path)
    
    def load_sheet_list(self, data_file):
        """加载数据文件的工作表列表或处理CSV文件"""
        try:
            # 清除工作表下拉框
            self.sheet_combo.clear()
            
            if self.is_csv_file:
                # 对于CSV文件，没有工作表的概念
                self.sheet_combo.addItem("CSV数据")
                self.sheet_combo.setEnabled(False)  # 禁用工作表选择
            else:
                # 使用pandas读取Excel工作表列表
                excel_data = pd.ExcelFile(data_file)
                sheet_names = excel_data.sheet_names
                
                self.sheet_combo.addItems(sheet_names)
                self.sheet_combo.setEnabled(True)  # 启用工作表选择
        except Exception as e:
            QMessageBox.warning(self, "错误", f"无法读取数据文件: {str(e)}")
    
    def log_message(self, message):
        """记录操作日志"""
        print(f"[Excel处理] {message}")
    
    def extract_excel_data(self):
        """提取Excel或CSV数据"""
        data_file = self.excel_input.text().strip()
        
        if not data_file:
            QMessageBox.warning(self, "错误", "请选择数据文件")
            return
        
        if not os.path.exists(data_file):
            QMessageBox.warning(self, "错误", "数据文件不存在")
            return
        
        # 获取选项
        sheet_name = self.sheet_combo.currentText()
        start_range = self.range_start.text().strip()
        end_range = self.range_end.text().strip()
        mode = self.mode_combo.currentText()
        
        try:
            # 根据文件类型选择不同的读取方法
            if self.is_csv_file:
                # 读取CSV文件
                # 尝试不同的编码，解决中文乱码问题
                try:
                    df = pd.read_csv(data_file, encoding='utf-8')
                except UnicodeDecodeError:
                    try:
                        df = pd.read_csv(data_file, encoding='gbk')
                    except:
                        df = pd.read_csv(data_file, encoding='latin1')
            else:
                # 读取Excel文件
                df = pd.read_excel(data_file, sheet_name=sheet_name)
            
            # 保存原始数据用于范围选择
            original_df = df.copy()
            
            # 处理提取范围
            if start_range and end_range:
                # 实现根据单元格范围提取数据的逻辑
                try:
                    # 提取起始和结束单元格的行和列
                    start_col = ''.join(filter(str.isalpha, start_range))
                    start_row = int(''.join(filter(str.isdigit, start_range)))
                    end_col = ''.join(filter(str.isalpha, end_range))
                    end_row = int(''.join(filter(str.isdigit, end_range)))
                    
                    # 获取列索引
                    cols = list(df.columns)
                    # 尝试直接匹配列名
                    if start_col in cols and end_col in cols:
                        start_col_idx = cols.index(start_col)
                        end_col_idx = cols.index(end_col)
                    else:
                        # 如果找不到列名，尝试将字母转换为索引（A=0, B=1, 等）
                        start_col_idx = sum((ord(c.upper()) - ord('A') + 1) * (26 ** i) for i, c in enumerate(reversed(start_col)))
                        end_col_idx = sum((ord(c.upper()) - ord('A') + 1) * (26 ** i) for i, c in enumerate(reversed(end_col)))
                        start_col_idx -= 1  # 转换为0-index
                        end_col_idx -= 1
                        
                        # 确保索引在有效范围内
                        start_col_idx = max(0, min(start_col_idx, len(cols)-1))
                        end_col_idx = max(0, min(end_col_idx, len(cols)-1))
                    
                    # 选择数据范围
                    df = df.iloc[start_row-1:end_row, start_col_idx:end_col_idx+1]
                    QMessageBox.information(self, "成功", f"已提取范围 {start_range} 到 {end_range} 的数据")
                except Exception as e:
                    QMessageBox.warning(self, "警告", f"范围格式不正确或无法提取指定范围: {str(e)}")
                    df = original_df  # 恢复原始数据
            
            # 根据提取模式处理数据
            if mode == "仅特定列" and start_range and not end_range:
                # 仅提取特定列
                if start_range in df.columns:
                    df = df[[start_range]]
                    QMessageBox.information(self, "成功", f"已提取列 '{start_range}' 的数据")
                else:
                    # 尝试将输入解析为列索引
                    try:
                        # 支持A1格式（只取列部分）或纯数字索引
                        if start_range.isdigit():
                            col_idx = int(start_range) - 1
                        else:
                            # 解析A1格式中的列名（如A, B, C...）
                            col_letters = ''.join(filter(str.isalpha, start_range))
                            col_idx = sum((ord(c.upper()) - ord('A') + 1) * (26 ** i) for i, c in enumerate(reversed(col_letters)))
                            col_idx -= 1  # 转换为0-index
                        
                        if 0 <= col_idx < len(df.columns):
                            col_name = df.columns[col_idx]
                            df = df[[col_name]]
                            QMessageBox.information(self, "成功", f"已提取第 {col_idx+1} 列（列名: {col_name}）的数据")
                        else:
                            QMessageBox.warning(self, "警告", f"列索引 {col_idx+1} 超出范围（总列数: {len(df.columns)}）")
                    except:
                        QMessageBox.warning(self, "警告", f"找不到列 '{start_range}'，将显示所有数据")
            elif mode == "仅特定行" and start_range and not end_range:
                # 仅提取特定行
                try:
                    # 尝试解析为行号
                    row_idx = int(start_range) - 1
                    if 0 <= row_idx < len(df):
                        df = df.iloc[[row_idx]]
                        QMessageBox.information(self, "成功", f"已提取第 {start_range} 行的数据")
                    else:
                        QMessageBox.warning(self, "警告", f"行号 {start_range} 超出范围（总行数: {len(df)}）")
                except:
                    QMessageBox.warning(self, "警告", f"行号格式不正确，将显示所有数据")
            
            # 显示数据到表格
            self.result_table.setRowCount(len(df))
            self.result_table.setColumnCount(len(df.columns))
            
            # 设置表头
            self.result_table.setHorizontalHeaderLabels(df.columns)
            
            # 填充数据
            for row_idx, row in df.iterrows():
                for col_idx, value in enumerate(row):
                    # 处理NaN值
                    display_value = "" if pd.isna(value) else str(value)
                    self.result_table.setItem(row_idx, col_idx, QTableWidgetItem(display_value))
            
            if not (start_range and (end_range or (mode != "所有数据"))):
                QMessageBox.information(self, "成功", "数据提取完成")
            
        except Exception as e:
            QMessageBox.warning(self, "错误", f"提取数据时出错: {str(e)}")
    
    def export_to_txt(self):
        """导出为TXT文件"""
        self.log_message("开始导出TXT文件")
        
        if self.result_table.rowCount() == 0:
            QMessageBox.warning(self, "错误", "没有数据可导出")
            self.log_message("导出失败: 没有数据可导出")
            return
        
        # 获取用户选择的保存路径
        file_path, _ = QFileDialog.getSaveFileName(
            self, "导出TXT文件", "", "文本文件 (*.txt);;所有文件 (*)"
        )
        
        if not file_path:
            self.log_message("用户取消了导出操作")
            return  # 用户取消了保存
        
        try:
            # 确保文件有正确的扩展名
            if not file_path.endswith('.txt'):
                original_path = file_path
                file_path += '.txt'
                self.log_message(f"添加扩展名 .txt: {original_path} -> {file_path}")
            
            # 检查目录是否存在
            directory = os.path.dirname(file_path)
            if directory and not os.path.exists(directory):
                self.log_message(f"目标目录不存在，尝试创建: {directory}")
                os.makedirs(directory, exist_ok=True)
            
            # 获取表格数据
            headers = []
            data = []
            
            # 获取表头
            for col_idx in range(self.result_table.columnCount()):
                header_item = self.result_table.horizontalHeaderItem(col_idx)
                headers.append(header_item.text() if header_item else f"列{col_idx+1}")
            
            # 获取数据
            for row_idx in range(self.result_table.rowCount()):
                row_data = []
                for col_idx in range(self.result_table.columnCount()):
                    item = self.result_table.item(row_idx, col_idx)
                    row_data.append(item.text() if item else "")
                data.append(row_data)
            
            self.log_message(f"成功获取表格数据: {len(data)} 行, {len(headers)} 列")
            
            # 导出为TXT文件，使用制表符分隔
            with open(file_path, 'w', encoding='utf-8-sig') as f:
                # 写入表头
                f.write('\t'.join(headers) + '\n')
                # 写入数据
                for row in data:
                    f.write('\t'.join(row) + '\n')
                # 确保数据被写入磁盘
                f.flush()
                os.fsync(f.fileno())
            
            # 验证文件是否已创建并且不为空
            if os.path.exists(file_path) and os.path.getsize(file_path) > 0:
                self.log_message(f"数据成功导出到: {file_path}, 文件大小: {os.path.getsize(file_path)} 字节")
                QMessageBox.information(self, "成功", f"数据已导出到: {file_path}")
            else:
                self.log_message(f"导出失败: 文件创建失败或为空, 路径: {file_path}")
                QMessageBox.critical(self, "错误", f"文件创建失败或为空\n路径: {file_path}")
        except PermissionError as e:
            self.log_message(f"权限错误: 无法写入文件 {file_path}")
            QMessageBox.critical(self, "错误", f"无法写入文件: {file_path}\n\n请检查文件是否被其他程序占用或您是否有写入权限。\n\n权限错误详情: {str(e)}")
        except Exception as e:
            self.log_message(f"导出TXT时出错: {str(e)}")
            QMessageBox.critical(self, "错误", f"导出TXT时出错: {str(e)}")
    
    def export_to_csv(self):
        """导出为CSV"""
        self.log_message("开始导出CSV文件")
        
        if self.result_table.rowCount() == 0:
            QMessageBox.warning(self, "错误", "没有数据可导出")
            self.log_message("导出失败: 没有数据可导出")
            return
        
        file_path, _ = QFileDialog.getSaveFileName(
            self, "导出CSV文件", "", "CSV文件 (*.csv);;所有文件 (*)"
        )
        
        if not file_path:
            self.log_message("用户取消了导出操作")
            return  # 用户取消了保存
        
        try:
            # 确保文件有正确的扩展名
            if not file_path.endswith('.csv'):
                original_path = file_path
                file_path += '.csv'
                self.log_message(f"添加扩展名 .csv: {original_path} -> {file_path}")
            
            # 检查目录是否存在
            directory = os.path.dirname(file_path)
            if directory and not os.path.exists(directory):
                self.log_message(f"目标目录不存在，尝试创建: {directory}")
                os.makedirs(directory, exist_ok=True)
            
            # 获取表格数据
            headers = []
            data = []
            
            # 获取表头
            for col_idx in range(self.result_table.columnCount()):
                header_item = self.result_table.horizontalHeaderItem(col_idx)
                headers.append(header_item.text() if header_item else f"列{col_idx+1}")
            
            # 获取数据
            for row_idx in range(self.result_table.rowCount()):
                row_data = []
                for col_idx in range(self.result_table.columnCount()):
                    item = self.result_table.item(row_idx, col_idx)
                    row_data.append(item.text() if item else "")
                data.append(row_data)
            
            # 创建DataFrame并导出为CSV
            df = pd.DataFrame(data, columns=headers)
            df.to_csv(file_path, index=False, encoding='utf-8-sig')
            
            # 验证文件是否已创建并且不为空
            if os.path.exists(file_path) and os.path.getsize(file_path) > 0:
                self.log_message(f"数据成功导出到: {file_path}, 文件大小: {os.path.getsize(file_path)} 字节")
                QMessageBox.information(self, "成功", f"数据已导出到: {file_path}")
            else:
                self.log_message(f"导出失败: 文件创建失败或为空, 路径: {file_path}")
                QMessageBox.critical(self, "错误", f"文件创建失败或为空\n路径: {file_path}")
        except PermissionError as e:
            self.log_message(f"权限错误: 无法写入文件 {file_path}")
            QMessageBox.critical(self, "错误", f"无法写入文件: {file_path}\n\n请检查文件是否被其他程序占用或您是否有写入权限。\n\n权限错误详情: {str(e)}")
        except Exception as e:
            self.log_message(f"导出CSV时出错: {str(e)}")
            QMessageBox.critical(self, "错误", f"导出CSV时出错: {str(e)}")
    
    def export_to_excel(self):
        """导出为Excel"""
        self.log_message("开始导出Excel文件")
        
        if self.result_table.rowCount() == 0:
            QMessageBox.warning(self, "错误", "没有数据可导出")
            self.log_message("导出失败: 没有数据可导出")
            return
        
        file_path, _ = QFileDialog.getSaveFileName(
            self, "导出Excel文件", "", "Excel文件 (*.xlsx);;所有文件 (*)"
        )
        
        if not file_path:
            self.log_message("用户取消了导出操作")
            return  # 用户取消了保存
        
        try:
            # 确保文件有正确的扩展名
            if not file_path.endswith('.xlsx'):
                original_path = file_path
                file_path += '.xlsx'
                self.log_message(f"添加扩展名 .xlsx: {original_path} -> {file_path}")
            
            # 检查目录是否存在
            directory = os.path.dirname(file_path)
            if directory and not os.path.exists(directory):
                self.log_message(f"目标目录不存在，尝试创建: {directory}")
                os.makedirs(directory, exist_ok=True)
            
            # 获取表格数据
            headers = []
            data = []
            
            # 获取表头
            for col_idx in range(self.result_table.columnCount()):
                header_item = self.result_table.horizontalHeaderItem(col_idx)
                headers.append(header_item.text() if header_item else f"列{col_idx+1}")
            
            # 获取数据
            for row_idx in range(self.result_table.rowCount()):
                row_data = []
                for col_idx in range(self.result_table.columnCount()):
                    item = self.result_table.item(row_idx, col_idx)
                    row_data.append(item.text() if item else "")
                data.append(row_data)
            
            # 创建DataFrame并导出为Excel
            df = pd.DataFrame(data, columns=headers)
            df.to_excel(file_path, index=False)
            
            # 验证文件是否已创建并且不为空
            if os.path.exists(file_path) and os.path.getsize(file_path) > 0:
                self.log_message(f"数据成功导出到: {file_path}, 文件大小: {os.path.getsize(file_path)} 字节")
                QMessageBox.information(self, "成功", f"数据已导出到: {file_path}")
            else:
                self.log_message(f"导出失败: 文件创建失败或为空, 路径: {file_path}")
                QMessageBox.critical(self, "错误", f"文件创建失败或为空\n路径: {file_path}")
        except PermissionError as e:
            self.log_message(f"权限错误: 无法写入文件 {file_path}")
            QMessageBox.critical(self, "错误", f"无法写入文件: {file_path}\n\n请检查文件是否被其他程序占用或您是否有写入权限。\n\n权限错误详情: {str(e)}")
        except Exception as e:
            self.log_message(f"导出Excel时出错: {str(e)}")
            QMessageBox.critical(self, "错误", f"导出Excel时出错: {str(e)}")


# 添加Word转换线程类
class WordToPDFThread(QThread):
    finished = pyqtSignal(bool, str)  # 成功标志和消息
    progress_updated = pyqtSignal(int)  # 进度更新信号
    
    # 类变量，用于存储共享的Word应用程序实例
    _shared_word_app = None
    _word_app_lock = QMutex()
    
    @classmethod
    def get_word_app(cls):
        with QMutexLocker(cls._word_app_lock):
            if cls._shared_word_app is None:
                try:
                    cls._shared_word_app = comtypes.client.CreateObject('Word.Application')
                    cls._shared_word_app.Visible = False
                except Exception as e:
                    print(f"创建Word应用程序失败: {str(e)}")
                    return None
            return cls._shared_word_app
    
    @classmethod
    def release_word_app(cls):
        with QMutexLocker(cls._word_app_lock):
            if cls._shared_word_app is not None:
                try:
                    cls._shared_word_app.Quit()
                except:
                    pass
                cls._shared_word_app = None
    
    def __init__(self, word_path, output_path, keep_layout, optimize, fast_mode=False):
        super().__init__()
        self.word_path = word_path
        self.output_path = output_path
        self.keep_layout = keep_layout
        self.optimize = optimize
        self.fast_mode = fast_mode
        self._canceled = False
    
    def cancel(self):
        """取消转换"""
        self._canceled = True
    
    def run(self):
        try:
            # 在Windows上使用COM接口转换Word到PDF
            if os.name != 'nt':  # 非Windows系统
                self.finished.emit(False, "此功能仅在Windows系统上可用")
                return
            
            # 获取共享的Word应用程序实例
            word = self.get_word_app()
            if word is None:
                self.finished.emit(False, "无法创建或获取Word应用程序实例")
                return
            
            try:
                # 打开文档
                self.progress_updated.emit(20)  # 报告进度
                
                if self.fast_mode:
                    # 快速模式下的打开选项
                    doc = word.Documents.Open(
                        os.path.abspath(self.word_path),
                        ReadOnly=True,  # 只读模式打开
                        AddToRecentFiles=False,  # 不添加到最近文件
                        Visible=False  # 不可见
                    )
                else:
                    doc = word.Documents.Open(os.path.abspath(self.word_path))
                
                if self._canceled:
                    doc.Close(SaveChanges=0)  # 不保存关闭
                    self.finished.emit(False, "转换已取消")
                    return
                
                self.progress_updated.emit(50)  # 报告进度
                
                # 设置PDF导出选项
                pdf_format = 17  # PDF格式
                export_options = {}
                
                if self.optimize:
                    # 优化文件大小
                    export_options['OptimizeFor'] = 1  # 最小文件大小
                
                if not self.keep_layout:
                    # 不保持严格布局，提高速度
                    export_options['UseISO19005_1'] = True  # 使用ISO标准提高兼容性
                
                # 导出为PDF
                self.progress_updated.emit(70)  # 报告进度
                
                if self.fast_mode:
                    # 快速模式下简化导出选项
                    doc.SaveAs(os.path.abspath(self.output_path), FileFormat=pdf_format)
                else:
                    # 完整模式，使用所有选项
                    doc.ExportAsFixedFormat(
                        OutputFileName=os.path.abspath(self.output_path),
                        ExportFormat=pdf_format,
                        OpenAfterExport=False,
                        OptimizeFor=1 if self.optimize else 0,
                        CreateBookmarks=0,  # 不创建书签
                        IncludeDocProps=False,
                        KeepIRM=False,
                        DocStructureTags=False,
                        BitmapMissingFonts=False,
                        UseISO19005_1=not self.keep_layout
                    )
                
                self.progress_updated.emit(90)  # 报告进度
                doc.Close(SaveChanges=0)  # 不保存关闭
                
                if not self._canceled:
                    self.progress_updated.emit(100)
                    self.finished.emit(True, self.output_path)
            except Exception as e:
                self.finished.emit(False, str(e))
        except Exception as e:
            self.finished.emit(False, str(e))


# 添加PDF转换线程类
class PDFToWordThread(QThread):
    finished = pyqtSignal(bool, str)  # 成功标志和消息
    
    def __init__(self, pdf_path, output_path, extract_images):
        super().__init__()
        self.pdf_path = pdf_path
        self.output_path = output_path
        self.extract_images = extract_images
    
    def run(self):
        try:
            # 创建Word文档
            doc = Document()
            
            # 打开PDF文件
            pdf_document = fitz.open(self.pdf_path)
            total_pages = len(pdf_document)
            
            # 为保存图片创建临时目录
            image_dir = None
            if self.extract_images:
                image_dir = os.path.join(os.path.dirname(self.output_path), 
                                        f"{os.path.basename(self.output_path)}_images")
                os.makedirs(image_dir, exist_ok=True)
            
            # 遍历每一页
            for page_num in range(total_pages):
                page = pdf_document.load_page(page_num)
                
                # 提取文本
                text = page.get_text()
                if text.strip():
                    doc.add_heading(f'第{page_num+1}页', level=2)
                    doc.add_paragraph(text)
                
                # 提取图片
                if self.extract_images:
                    image_list = page.get_images(full=True)
                    for img_index, img in enumerate(image_list):
                        xref = img[0]
                        base_image = pdf_document.extract_image(xref)
                        image_bytes = base_image["image"]
                        image_ext = base_image["ext"]
                        
                        # 保存图片
                        image_filename = f"page_{page_num+1}_img_{img_index+1}.{image_ext}"
                        image_path = os.path.join(image_dir, image_filename)
                        
                        with open(image_path, "wb") as f:
                            f.write(image_bytes)
                        
                        # 添加图片到Word文档
                        doc.add_paragraph(f"图片 {img_index+1}:")
                        doc.add_picture(image_path, width=Inches(6))
                
                # 添加分页符
                if page_num < total_pages - 1:
                    doc.add_page_break()
            
            # 保存Word文档
            doc.save(self.output_path)
            pdf_document.close()
            
            self.finished.emit(True, self.output_path)
        except Exception as e:
            self.finished.emit(False, str(e))